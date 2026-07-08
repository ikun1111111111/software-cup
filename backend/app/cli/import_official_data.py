"""Import official Lingshan scenic area DOCX materials into PG+Milvus.

Usage:
    python -m app.cli.import_official_data \
        --guide ".../灵山胜境：历史、文化、景点特色与个性化游览指南.docx" \
        --spots ".../灵山胜境 景点结构化数据集.docx"
"""
import argparse
import asyncio
import json
import logging
import os
import re
from datetime import datetime
from typing import Any

from docx import Document
from sqlalchemy import select
from sqlalchemy.dialects.postgresql import insert

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("import_official")


# ---------------------------------------------------------------------------
# Topic classification for guide headings
# ---------------------------------------------------------------------------

def classify_heading_topic(heading: str) -> str:
    """Map a section heading to a knowledge topic."""
    h = heading.lower()
    if "历史" in h or "渊源" in h or "兴衰" in h:
        return "history"
    if "文化" in h:
        return "culture"
    if "路线" in h or "攻略" in h or "游览" in h and ("推荐" in h or "路线" in h):
        return "route"
    if "门票" in h or "票价" in h or "优惠" in h or "价格" in h:
        return "ticket"
    if "餐饮" in h or "美食" in h or "住宿" in h or "素斋" in h or "素面" in h:
        return "food"
    return "general"


# ---------------------------------------------------------------------------
# DOCX parsing helpers
# ---------------------------------------------------------------------------

def iter_paragraphs(doc: Document):
    """Yield (style_name, text) for each non-empty paragraph."""
    for p in doc.paragraphs:
        style = p.style.name if p.style else "Normal"
        text = p.text.strip()
        if text:
            yield style, text


def parse_guide_sections(doc: Document) -> list[dict[str, Any]]:
    """Split guide DOCX into top-level sections by Heading 1.

    Returns list of {"title", "topic", "text"}.
    """
    sections = []
    current_title = "景区概况"
    current_topic = "general"
    current_lines = []

    for style, text in iter_paragraphs(doc):
        if style.startswith("Heading 1"):
            if current_lines:
                sections.append({
                    "title": current_title,
                    "topic": current_topic,
                    "text": "\n".join(current_lines),
                })
            current_title = text
            current_topic = classify_heading_topic(text)
            current_lines = []
        elif style.startswith("Heading"):
            current_lines.append(f"【{text}】")
        else:
            current_lines.append(text)

    if current_lines:
        sections.append({
            "title": current_title,
            "topic": current_topic,
            "text": "\n".join(current_lines),
        })

    return sections


def parse_spots_tables(doc: Document) -> list[dict[str, str]]:
    """Parse scenic spot structured tables into dict records."""
    records = []
    headers = [
        "景区名称", "景点ID", "景点名称", "具体位置", "建筑/景观参数",
        "核心功能", "文化内涵", "详细介绍", "游玩亮点", "演艺/开放信息", "备注",
    ]
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if not cells or not any(cells):
                continue
            rows.append(cells)
        if not rows:
            continue
        # Detect header row
        first = rows[0]
        if "景点ID" in first or "景区名称" in first:
            header = first
            data_rows = rows[1:]
        else:
            header = headers[: len(first)]
            data_rows = rows

        for cells in data_rows:
            row_dict = {}
            for i, key in enumerate(header):
                if i < len(cells):
                    row_dict[key] = cells[i]
            if row_dict.get("景点ID") or row_dict.get("景点名称"):
                records.append(row_dict)
    return records


def extract_ticket_faq(text: str) -> list[dict[str, str]]:
    """Extract ticket-related FAQ entries from guide text and static facts."""
    faqs = []
    patterns = [
        (r"成人票[\s|]*([0-9]+)\s*元", "灵山胜境成人票多少钱？", "灵山胜境成人票为{price}元。"),
        (r"半价票[\s|]*([0-9]+)\s*元", "灵山胜境半价票多少钱？", "灵山胜境半价票为{price}元，适用于6-18周岁未成年人、全日制本科及以下学生、60-69周岁老人。"),
        (r"免票[\s|]*0\s*元", "灵山胜境哪些人免票？", "6周岁以下或1.4米以下儿童、70周岁以上老人、现役军人、残疾人可免票入园。"),
        (r"《灵山吉祥颂》.*?([0-9]{1,2}:[0-9]{2}).*?([0-9]{1,2}:[0-9]{2})", "灵山吉祥颂演出时间？", "《灵山吉祥颂》演出时间一般为{time1}、{time2}等，具体以景区公告为准，每场时长约20分钟。"),
        (r"九龙灌浴.*?([0-9]{1,2}:[0-9]{2}).*?([0-9]{1,2}:[0-9]{2})", "九龙灌浴演出时间？", "九龙灌浴平日演出时间为{time1}、{time2}等，周末及节假日会增加场次，具体以景区广播通知为准。"),
    ]
    for pattern, question, answer_template in patterns:
        m = re.search(pattern, text)
        if m:
            if "{price}" in answer_template:
                answer = answer_template.format(price=m.group(1))
            elif "{time1}" in answer_template:
                answer = answer_template.format(time1=m.group(1), time2=m.group(2))
            else:
                answer = answer_template
            faqs.append({"question": question, "answer": answer, "category": "ticket"})
    return faqs


def build_static_ticket_faqs() -> list[dict[str, str]]:
    """Return static FAQ facts covering all six knowledge topics."""
    return [
        # ticket
        {"question": "灵山胜境门票多少钱？", "answer": "灵山胜境成人票约210元/人；半价票约105元，适用于6-18周岁未成年人、全日制本科及以下学生、60-69周岁老人；6周岁以下或1.4米以下儿童、70周岁以上老人、现役军人、残疾人免票。", "category": "ticket", "keywords": "门票,票价,多少钱,价格"},
        {"question": "灵山胜境开放时间是什么时候？", "answer": "灵山胜境一般开放时间为8:00-17:00，冬季闭园时间可能提前至16:30，具体以景区公告为准。", "category": "ticket", "keywords": "开放时间,几点开门,几点关门,营业时间"},
        {"question": "灵山胜境有观光车吗？", "answer": "景区提供观光车服务，单独购票40元/人，也可购买门票+观光车联票225元，无限次乘坐更划算。", "category": "ticket", "keywords": "观光车,电瓶车,游览车,交通"},
        {"question": "《灵山吉祥颂》演出时间是什么？", "answer": "《灵山吉祥颂》演出时间一般为每日10:35、11:30、14:00、16:00，每场时长约20分钟，节假日可能加演，具体以景区公告为准。", "category": "ticket", "keywords": "吉祥颂,演出,时间,场次"},
        {"question": "九龙灌浴演出时间是什么？", "answer": "九龙灌浴平日演出时间一般为10:00、11:30、13:30、15:00，周末及节假日会增加场次，每场约15分钟，具体以景区广播通知为准。", "category": "ticket", "keywords": "九龙灌浴,演出,时间,场次"},
        {"question": "灵山胜境半价票多少钱？", "answer": "灵山胜境半价票约105元，适用于6-18周岁未成年人、全日制本科及以下学生、60-69周岁老人。", "category": "ticket", "keywords": "半价票,优惠,学生,老人"},
        {"question": "灵山胜境哪些人免票？", "answer": "6周岁以下或1.4米以下儿童、70周岁以上老人、现役军人、残疾人可免票入园。", "category": "ticket", "keywords": "免票,免费,儿童,老人,军人,残疾人"},
        {"question": "灵山胜境门票加观光车联票多少钱？", "answer": "门票+观光车联票约225元，可无限次乘坐观光车。", "category": "ticket", "keywords": "联票,观光车,多少钱,套票"},
        {"question": "灵山胜境儿童票多少钱？", "answer": "6-18周岁未成年人可购买半价票约105元；6周岁以下或1.4米以下儿童免票。", "category": "ticket", "keywords": "儿童票,小孩,多少钱,免票"},
        {"question": "灵山胜境老人票有优惠吗？", "answer": "60-69周岁老人可购买半价票约105元，70周岁以上老人免票。", "category": "ticket", "keywords": "老人票,老年证,优惠,免票"},
        # food
        {"question": "灵山胜境有什么特色美食？", "answer": "景区以素斋、素面为特色，梵宫素斋自助约50元/位，素面套餐约35元/位，灵山精舍也提供精致素斋。", "category": "food", "keywords": "美食,特色,素斋,素面,吃什么"},
        {"question": "梵宫素斋自助多少钱？", "answer": "梵宫素斋自助约50元/位，菜品丰富、清淡雅致。", "category": "food", "keywords": "素斋,自助,多少钱,梵宫"},
        {"question": "灵山素面套餐多少钱？", "answer": "素面套餐约35元/位，景区内多个餐厅提供。", "category": "food", "keywords": "素面,套餐,多少钱"},
        {"question": "灵山精舍有素斋吗？", "answer": "有，灵山精舍提供禅意素斋，环境优雅、菜品精致。", "category": "food", "keywords": "灵山精舍,素斋,餐厅"},
        {"question": "灵山胜境附近有什么好吃的？", "answer": "景区及周边以素斋、素面、江南小吃为主，可在梵宫、灵山精舍及马山镇品尝。", "category": "food", "keywords": "附近,好吃的,餐厅,素斋"},
        # route
        {"question": "推荐一条灵山胜境历史文化路线？", "answer": "历史文化爱好者路线（约6小时）：南门入园→灵山大照壁→胜境广场→佛手广场→祥符禅寺→杏坛广场→佛前广场→灵山大佛→灵山梵宫→五印坛城→三圣殿→出口。重点了解玄奘与小灵山渊源、祥符禅寺千年兴衰、梵宫非遗艺术等。", "category": "route", "keywords": "历史文化,路线,推荐,游览"},
        {"question": "推荐一条灵山胜境亲子路线？", "answer": "亲子家庭路线（约4小时）：南门入园→九龙灌浴→佛手广场→百子戏弥勒→梵宫→五印坛城→出口。可参与抱佛脚、摸佛手、看吉祥颂等互动体验。", "category": "route", "keywords": "亲子,路线,推荐,孩子"},
        {"question": "推荐一条灵山胜境自然风光路线？", "answer": "自然风光爱好者路线（约5小时）：南门入园→佛足坛→九龙灌浴→菩提大道→灵山大佛→曼飞龙塔→灵山精舍→梵宫广场→出口。可欣赏太湖风光、菩提林荫和禅意园林。", "category": "route", "keywords": "自然风光,路线,推荐,全景"},
        {"question": "游览灵山胜境大概需要多长时间？", "answer": "一般游览需半天到一天。历史文化深度游约6小时，自然风光全景游约5小时，亲子轻松游约4小时。", "category": "route", "keywords": "多长时间,多久,游览,小时"},
        {"question": "灵山胜境的最佳游览时间是什么时候？", "answer": "春秋季节（3-5月、9-11月）气候宜人，适合户外活动；建议上午9点前入园避开人流高峰。", "category": "route", "keywords": "最佳游览时间,季节,什么时候去"},
        {"question": "第一次去灵山胜境应该怎么游览？", "answer": "建议上午先参观灵山大佛和祥符禅寺，中午观看九龙灌浴演出，下午游览梵宫及吉祥颂演出，傍晚可前往拈花湾欣赏夜景。", "category": "route", "keywords": "第一次,怎么游览,攻略"},
        {"question": "灵山胜境一日游怎么安排？", "answer": "一日游推荐：上午灵山大佛、祥符禅寺；中午九龙灌浴；下午梵宫、吉祥颂；傍晚拈花湾夜景。", "category": "route", "keywords": "一日游,安排,攻略"},
        # history
        {"question": "小灵山这个名字是怎么来的？", "answer": "唐贞观年间，玄奘法师西行取经归来，途经马山时见此地山形酷似印度灵鹫山，遂将“灵鹫胜境”之名赐予此地，命名为“小灵山”，并嘱弟子窥基法师在此住持道场。", "category": "history", "keywords": "小灵山,名字,由来,玄奘"},
        {"question": "祥符禅寺有什么历史？", "answer": "祥符禅寺始建于唐贞观年间，由玄奘法师弟子窥基法师开坛讲经；北宋大中祥符年间（1008-1016年）宋真宗赐额“祥符禅寺”，历经兴废，是江南千年禅宗祖庭。", "category": "history", "keywords": "祥符禅寺,历史,唐,北宋"},
        {"question": "灵山大佛是什么时候建成开光的？", "answer": "灵山大佛于1994年奠基，1997年11月15日落成开光，成为景区标志性建筑。", "category": "history", "keywords": "灵山大佛,建成,开光,1997"},
        {"question": "灵山胜境和唐玄奘有什么关系？", "answer": "玄奘法师西行归来后命名小灵山，并嘱弟子窥基法师在此住持道场，奠定了灵山佛教根基。", "category": "history", "keywords": "玄奘,关系,小灵山"},
        {"question": "灵山胜境的建设历程是怎样的？", "answer": "1994年奠基修复祥符禅寺、建造灵山大佛；1997年灵山大佛开光；2003年建成九龙灌浴；2006-2009年建成梵宫、五印坛城、曼飞龙塔等，形成综合性佛教文化景区。", "category": "history", "keywords": "建设,历程,发展"},
        {"question": "灵山胜境获得过什么荣誉？", "answer": "灵山胜境是国家5A级旅游景区、世界佛教论坛永久会址，被誉为“东方佛国”和“太湖佛国”。", "category": "history", "keywords": "荣誉,5A,世界佛教论坛"},
        {"question": "灵山大佛有多高？", "answer": "灵山大佛通高88米（佛体79米，莲花瓣9米），含台基总高101.5米，总用铜量725吨，是世界最高的露天青铜释迦牟尼立像。", "category": "history", "keywords": "灵山大佛,多高,88米"},
        {"question": "灵山大佛铸造用了多少吨铜？", "answer": "灵山大佛总用铜量约725吨，由2000块铸铜面板拼接而成。", "category": "history", "keywords": "铜,吨,铸造"},
        {"question": "灵山胜境的历史可以追溯到什么时候？", "answer": "灵山胜境的历史可追溯至1300多年前的唐代贞观年间，与玄奘法师西行取经紧密相连。", "category": "history", "keywords": "历史,追溯,唐代"},
        {"question": "赵朴初和灵山胜境有什么关系？", "answer": "赵朴初先生提出“五方五佛”之论，认为东方空缺，支持建造灵山大佛；他还为景区题写“灵山胜境”。", "category": "history", "keywords": "赵朴初,五方五佛,题字"},
        # culture
        {"question": "灵山梵宫有什么特色？", "answer": "灵山梵宫建筑面积7.2万平方米，造价18亿，被誉为“东方卢浮宫”和“佛教艺术的卢浮宫”，是世界佛教论坛主会场，内部汇集东阳木雕、琉璃、油画、景泰蓝、玉雕、漆画等非遗艺术精品。", "category": "culture", "keywords": "梵宫,特色,东方卢浮宫"},
        {"question": "九龙灌浴表演内容是什么？", "answer": "九龙灌浴依据释迦牟尼诞生传说打造，莲花绽放、太子佛升起，九条飞龙喷水沐浴太子佛，再现“花开见佛，九龙沐浴”的祥瑞场景。", "category": "culture", "keywords": "九龙灌浴,内容,释迦牟尼"},
        {"question": "五印坛城有什么文化内涵？", "answer": "“五印”代表五方五佛的五种手印，坛城即藏传佛教的曼陀罗道场，象征宇宙和谐圆满，体现汉藏佛教文化交融。", "category": "culture", "keywords": "五印坛城,五印,藏传佛教"},
        {"question": "曼飞龙塔是什么？", "answer": "曼飞龙塔复刻云南西双版纳曼飞龙白塔，为南传佛教干栏式建筑，九塔组合象征南传佛教的九种智慧与佛陀九种功德。", "category": "culture", "keywords": "曼飞龙塔,南传佛教"},
        {"question": "灵山胜境的佛教文化主要体现在哪些方面？", "answer": "主要体现在佛教建筑、雕塑、仪式、艺术与祈福体验中，如灵山大佛、梵宫非遗艺术、九龙灌浴、五印坛城、转经筒、撞钟祈福等。", "category": "culture", "keywords": "佛教文化,体现,建筑,艺术"},
        {"question": "灵山胜境有哪些祈福体验？", "answer": "主要祈福体验有九龙灌浴接圣水、天下第一掌摸佛手、抱佛脚、祥符禅寺撞钟、五印坛城转经筒等。", "category": "culture", "keywords": "祈福,体验,抱佛脚,转经筒"},
        {"question": "灵山梵宫有哪些非遗艺术？", "answer": "内部汇集东阳木雕、琉璃、油画、景泰蓝、玉雕、漆画等多种传统工艺精品。", "category": "culture", "keywords": "梵宫,非遗,木雕,琉璃"},
        {"question": "灵山胜境举办过什么重要活动？", "answer": "灵山胜境是世界佛教论坛永久会址，已成功举办多届世界佛教论坛。", "category": "culture", "keywords": "世界佛教论坛,活动"},
        {"question": "天下第一掌有什么寓意？", "answer": "“天下第一掌”为灵山大佛右手复制，高11.7米、宽5.5米，摸掌祈福寓意“沾福气、保平安”。", "category": "culture", "keywords": "天下第一掌,寓意,佛手"},
        {"question": "百子戏弥勒有什么寓意？", "answer": "百子戏弥勒体现佛教“皆大欢喜”理念，百子环绕寓意多子多福、家庭和睦、子孙满堂。", "category": "culture", "keywords": "百子戏弥勒,皆大欢喜"},
        # general
        {"question": "灵山胜境位于哪里？", "answer": "灵山胜境坐落于江苏省无锡市太湖西北部的马山镇，地处秦履峰、青龙山、白虎山三山环抱之间，是国家5A级旅游景区、世界佛教论坛永久会址。", "category": "general", "keywords": "位于,哪里,地址,无锡"},
        {"question": "灵山胜境有哪些主要景点？", "answer": "主要景点包括灵山大佛、灵山梵宫、九龙灌浴、五印坛城、祥符禅寺、曼飞龙塔、佛手广场、百子戏弥勒、灵山精舍等。", "category": "general", "keywords": "主要景点,有哪些"},
        {"question": "灵山大佛的佛手广场有什么特色？", "answer": "佛手广场有“天下第一掌”，为灵山大佛右手复制，高11.7米、宽5.5米，摸掌祈福寓意“沾福气、保平安”。", "category": "general", "keywords": "佛手广场,天下第一掌"},
        {"question": "降魔浮雕表现的是什么故事？", "answer": "降魔浮雕生动再现佛陀在菩提树下静坐修行，战胜魔王波旬诱惑与威胁，最终觉悟成佛的历程。", "category": "general", "keywords": "降魔浮雕,故事"},
        {"question": "阿育王柱有什么文化内涵？", "answer": "阿育王柱复刻古印度阿育王石柱造型，柱头四头狮子象征佛法向四方传播，彰显佛教和平、包容、普度精神。", "category": "general", "keywords": "阿育王柱,佛法传播"},
        {"question": "祥符禅寺在哪里？", "answer": "祥符禅寺位于灵山胜境中轴核心、灵山大佛基座之下，是景区内历史最悠久的人文景观。", "category": "general", "keywords": "祥符禅寺,在哪里"},
        {"question": "五明桥代表什么？", "answer": "五明桥代表佛教五种核心智慧：声明、因明、内明、医方明、工巧明，寓意过桥开启智慧、走向觉悟。", "category": "general", "keywords": "五明桥,智慧"},
        {"question": "佛足坛有什么寓意？", "answer": "佛足坛复刻佛祖释迦牟尼真身脚印，象征“佛足所至，佛光普照”，代表佛的福德与智慧圆满。", "category": "general", "keywords": "佛足坛,寓意"},
        {"question": "灵山胜境的导游服务多少钱？", "answer": "景区提供导游讲解服务，约300元起，适合希望深入了解景区历史文化的游客。", "category": "general", "keywords": "导游,讲解,多少钱"},
        {"question": "灵山胜境可以带宠物吗？", "answer": "景区为佛教文化场所，一般不允许携带宠物入园，请遵守景区规定、文明游览。", "category": "general", "keywords": "宠物,可以带吗"},
    ]


def extract_route_faqs(sections: list[dict[str, Any]]) -> list[dict[str, str]]:
    """Extract route-related FAQ from route sections."""
    faqs = []
    for sec in sections:
        if sec["topic"] != "route":
            continue
        # Each H2 inside route section is a route title
        for line in sec["text"].split("\n"):
            if line.startswith("【") and "路线" in line:
                title = line.strip("【】")
                faqs.append({
                    "question": f"推荐一条{title.replace('路线规划', '').replace('路线', '').strip()}？",
                    "answer": f"{title}详情请参考路线推荐资料。",
                    "category": "route",
                })
    return faqs


def extract_spot_faqs(records: list[dict[str, str]]) -> list[dict[str, str]]:
    """Generate FAQ entries for core spots."""
    faqs = []
    for r in records:
        name = r.get("景点名称", "")
        params = r.get("建筑/景观参数", "")
        if not name:
            continue
        # Height/parameter FAQ
        if "高" in params or "米" in params:
            faqs.append({
                "question": f"{name}有多高？",
                "answer": f"{name}相关参数：{params}",
                "category": "general",
            })
        # Location FAQ
        location = r.get("具体位置", "")
        if location:
            faqs.append({
                "question": f"{name}在哪里？",
                "answer": f"{name}位于{location}",
                "category": "general",
            })
    return faqs


# ---------------------------------------------------------------------------
# Database import
# ---------------------------------------------------------------------------

async def import_guide_sections(db, sections: list[dict[str, Any]], source_name: str):
    """Create KnowledgeDoc rows and trigger chunk/index processing."""
    from app.models.knowledge import KnowledgeDoc, DocStatus
    from app.services.knowledge_service import process_document

    imported = 0
    for sec in sections:
        if len(sec["text"]) < 50:
            continue
        # Upsert by title
        stmt = select(KnowledgeDoc).where(
            KnowledgeDoc.title == sec["title"],
        )
        existing = (await db.execute(stmt)).scalar_one_or_none()
        if existing:
            existing.content = sec["text"]
            existing.topic = sec["topic"]
            existing.topic_tags = json.dumps([sec["topic"]], ensure_ascii=False)
            existing.status = DocStatus.pending
            doc = existing
        else:
            doc = KnowledgeDoc(
                title=sec["title"],
                content=sec["text"],
                file_type="docx",
                status=DocStatus.pending,
                topic=sec["topic"],
                topic_tags=json.dumps([sec["topic"]], ensure_ascii=False),
            )
            db.add(doc)
        await db.commit()
        await db.refresh(doc)

        try:
            await process_document(doc.id, db)
            imported += 1
            logger.info("Indexed guide section: %s (topic=%s)", sec["title"], sec["topic"])
        except Exception as e:
            logger.warning("Failed to index guide section %s: %s", sec["title"], e)

    return imported


async def import_scenic_spots(db, records: list[dict[str, str]]):
    """Create or update ScenicSpot records from structured dataset."""
    from app.models.tourist import ScenicSpot

    imported = 0
    for r in records:
        spot_id = r.get("景点ID", "").strip()
        name = r.get("景点名称", "").strip()
        if not spot_id or not name:
            continue

        stmt = select(ScenicSpot).where(ScenicSpot.id == spot_id)
        existing = (await db.execute(stmt)).scalar_one_or_none()

        params = r.get("建筑/景观参数", "")
        topic_tags = []
        if "高" in params or "米" in params:
            topic_tags.append("general")
        if "佛教" in r.get("文化内涵", "") or "禅" in r.get("文化内涵", ""):
            topic_tags.append("culture")
        if "历史" in r.get("文化内涵", ""):
            topic_tags.append("history")

        values = {
            "id": spot_id,
            "name": name,
            "category": "核心景点" if r.get("景区名称") == "灵山胜境" else "拈花湾",
            "overview": r.get("核心功能", ""),
            "detail": r.get("详细介绍", ""),
            "tags": [r.get("文化内涵", ""), r.get("游玩亮点", "")],
            "topic_tags": topic_tags,
            "must_see": r.get("游玩亮点", ""),
            "open_time": r.get("演艺/开放信息", ""),
            "narration": r.get("详细介绍", ""),
            "is_active": True,
        }

        if existing:
            for k, v in values.items():
                setattr(existing, k, v)
            existing.updated_at = datetime.utcnow()
        else:
            values["created_at"] = datetime.utcnow()
            db.add(ScenicSpot(**values))

        imported += 1

    await db.commit()
    logger.info("Imported/updated %d scenic spots", imported)
    return imported


async def import_faqs(db, faqs: list[dict[str, str]]):
    """Upsert FAQ entries."""
    from app.models.knowledge import FaqEntry

    imported = 0
    for item in faqs:
        if not item.get("question") or not item.get("answer"):
            continue
        stmt = select(FaqEntry).where(FaqEntry.question == item["question"])
        existing = (await db.execute(stmt)).scalar_one_or_none()
        keywords = item.get("keywords")
        if existing:
            existing.answer = item["answer"]
            existing.category = item.get("category", "general")
            if keywords:
                existing.keywords = keywords
        else:
            db.add(FaqEntry(
                question=item["question"],
                answer=item["answer"],
                category=item.get("category", "general"),
                keywords=keywords,
                is_active=True,
            ))
        imported += 1
    await db.commit()
    logger.info("Imported/updated %d FAQ entries", imported)
    return imported


async def import_tour_routes(db, sections: list[dict[str, Any]]):
    """Create/update TourRoute records from route sections."""
    from app.models.tourist import TourRoute

    imported = 0
    for sec in sections:
        if sec["topic"] != "route":
            continue
        lines = sec["text"].split("\n")
        current_route = None
        current_description = []
        for line in lines:
            if line.startswith("【") and "路线" in line:
                # Save previous route
                if current_route:
                    await _upsert_route(db, current_route, "\n".join(current_description))
                    imported += 1
                title = line.strip("【】")
                route_id = "route_" + re.sub(r"[^\w]", "_", title)[:40]
                current_route = {
                    "id": route_id,
                    "name": title,
                    "duration": "约6小时" if "6小时" in title else ("约5小时" if "5小时" in title else ("约4小时" if "4小时" in title else "约5小时")),
                    "route_type": "culture" if "历史" in title else ("nature" if "自然" in title else "family"),
                }
                current_description = []
            elif current_route is not None:
                current_description.append(line)
        if current_route:
            await _upsert_route(db, current_route, "\n".join(current_description))
            imported += 1

    await db.commit()
    logger.info("Imported/updated %d tour routes", imported)
    return imported


async def _upsert_route(db, route: dict[str, Any], description: str):
    from app.models.tourist import TourRoute
    stmt = select(TourRoute).where(TourRoute.id == route["id"])
    existing = (await db.execute(stmt)).scalar_one_or_none()
    if existing:
        existing.name = route["name"]
        existing.duration = route["duration"]
        existing.route_type = route["route_type"]
        existing.description = description[:2000]
    else:
        db.add(TourRoute(
            id=route["id"],
            name=route["name"],
            duration=route["duration"],
            route_type=route["route_type"],
            description=description[:2000],
            spot_order=[],
        ))


# ---------------------------------------------------------------------------
# Main entry
# ---------------------------------------------------------------------------

def _find_material_dir() -> str | None:
    """Locate the official material directory by walking up from this file."""
    current = os.path.abspath(os.path.dirname(__file__))
    for _ in range(6):
        candidate = os.path.join(current, "软件杯a5官方素材")
        if os.path.isdir(candidate):
            return candidate
        parent = os.path.dirname(current)
        if parent == current:
            break
        current = parent
    return None


async def main_async():
    parser = argparse.ArgumentParser(description="Import official Lingshan DOCX data")
    parser.add_argument("--guide", default=None, help="Path to guide DOCX")
    parser.add_argument("--spots", default=None, help="Path to spots dataset DOCX")
    parser.add_argument("--skip-guide", action="store_true", help="Skip guide import")
    parser.add_argument("--skip-spots", action="store_true", help="Skip spots import")
    parser.add_argument("--skip-faqs", action="store_true", help="Skip FAQ import")
    parser.add_argument("--skip-routes", action="store_true", help="Skip route import")
    args = parser.parse_args()

    # Default paths relative to project root
    material_dir = _find_material_dir()
    default_guide = None
    default_spots = None
    if material_dir:
        default_guide = os.path.join(
            material_dir, "灵山胜境：历史、文化、景点特色与个性化游览指南.docx"
        )
        default_spots = os.path.join(
            material_dir, "灵山胜境 景点结构化数据集.docx"
        )

    guide_path = args.guide or default_guide
    spots_path = args.spots or default_spots

    from app.core.database import async_session

    async with async_session() as db:
        if not args.skip_guide and guide_path and os.path.exists(guide_path):
            logger.info("Parsing guide DOCX: %s", guide_path)
            guide_doc = Document(guide_path)
            sections = parse_guide_sections(guide_doc)
            logger.info("Found %d guide sections", len(sections))
            await import_guide_sections(db, sections, "official_guide")

            if not args.skip_routes:
                await import_tour_routes(db, sections)

            if not args.skip_faqs:
                all_text = "\n".join(s["text"] for s in sections)
                faqs = extract_ticket_faq(all_text)
                faqs.extend(build_static_ticket_faqs())
                faqs.extend(extract_route_faqs(sections))
                await import_faqs(db, faqs)

        if not args.skip_spots and spots_path and os.path.exists(spots_path):
            logger.info("Parsing spots dataset DOCX: %s", spots_path)
            spots_doc = Document(spots_path)
            records = parse_spots_tables(spots_doc)
            logger.info("Found %d spot records", len(records))
            await import_scenic_spots(db, records)

            if not args.skip_faqs:
                faqs = extract_spot_faqs(records)
                await import_faqs(db, faqs)

    logger.info("Import completed.")


def main():
    asyncio.run(main_async())


if __name__ == "__main__":
    main()
