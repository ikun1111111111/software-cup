"""Tests for document parsing service."""
import pytest
import os
from app.services.knowledge_service import parse_document


class TestParseService:
    """Test document parsing to plain text."""

    def test_parse_txt_file(self, tmp_path):
        """Should parse .txt file correctly."""
        file_path = tmp_path / "test.txt"
        content = "这是一段测试文本。\n包含多行内容。\n灵山胜境欢迎您！"
        file_path.write_text(content, encoding="utf-8")

        result = parse_document(str(file_path), "txt")

        assert result == content
        assert "灵山胜境" in result

    def test_parse_md_file(self, tmp_path):
        """Should parse .md markdown file."""
        file_path = tmp_path / "test.md"
        content = "# 灵山大佛\n\n灵山大佛高88米，是**中国最高**的青铜立佛。\n\n- 景点1\n- 景点2"
        file_path.write_text(content, encoding="utf-8")

        result = parse_document(str(file_path), "md")

        assert result == content
        assert "灵山大佛" in result
        assert "# " in result  # Raw markdown preserved

    def test_parse_docx_file(self, tmp_path):
        """Should parse .docx Word file."""
        from docx import Document

        file_path = tmp_path / "test.docx"
        doc = Document()
        doc.add_paragraph("第一段：灵山胜境简介。")
        doc.add_paragraph("第二段：主要景点介绍。")
        doc.save(str(file_path))

        result = parse_document(str(file_path), "docx")

        assert "第一段" in result
        assert "第二段" in result
        assert "灵山胜境简介" in result

    def test_parse_pdf_file(self, tmp_path):
        """Should parse .pdf file (requires PyMuPDF)."""
        import fitz

        file_path = tmp_path / "test.pdf"
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text(fitz.Point(72, 72), "灵山胜境测试PDF内容。", fontsize=12)
        doc.save(str(file_path))
        doc.close()

        result = parse_document(str(file_path), "pdf")

        assert "灵山胜境" in result
        assert "PDF" in result

    def test_parse_unsupported_type(self):
        """Should raise ValueError for unsupported file types."""
        with pytest.raises(ValueError, match="Unsupported"):
            parse_document("/fake/path.xyz", "xyz")

    def test_parse_empty_txt_file(self, tmp_path):
        """Should parse empty file without error."""
        file_path = tmp_path / "empty.txt"
        file_path.write_text("", encoding="utf-8")

        result = parse_document(str(file_path), "txt")
        assert result == ""

    def test_parse_multilingual_content(self, tmp_path):
        """Should handle mixed language content."""
        file_path = tmp_path / "mixed.txt"
        content = "灵山胜境 (Lingshan Scenic Area) 位于 Wuxi City.\nThe Big Buddha is 88 meters tall."
        file_path.write_text(content, encoding="utf-8")

        result = parse_document(str(file_path), "txt")

        assert "灵山胜境" in result
        assert "Lingshan" in result
        assert "88 meters" in result


class TestParseServiceEdgeCases:
    """Edge cases for document parsing."""

    def test_parse_file_with_special_chars(self, tmp_path):
        """Should handle special Unicode characters."""
        file_path = tmp_path / "special.txt"
        content = "测试内容：①②③ ★☆♠♣ ㊣㊧㊨ — 、。"
        file_path.write_text(content, encoding="utf-8")

        result = parse_document(str(file_path), "txt")

        assert "①②③" in result
        assert "★☆" in result
        assert "㊣" in result

    def test_parse_large_file(self, tmp_path):
        """Should handle reasonably large files."""
        file_path = tmp_path / "large.txt"
        content = "灵山胜境景区介绍。" * 10000  # ~50KB
        file_path.write_text(content, encoding="utf-8")

        result = parse_document(str(file_path), "txt")

        assert len(result) == len(content)
        assert result.count("灵山胜境") == 10000
