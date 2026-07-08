"""Tests for official DOCX import script."""
import pytest
from docx import Document

from app.cli.import_official_data import (
    classify_heading_topic,
    parse_guide_sections,
    parse_spots_tables,
    extract_ticket_faq,
    build_static_ticket_faqs,
    extract_route_faqs,
    extract_spot_faqs,
)


class TestTopicClassification:
    """Test heading-to-topic mapping."""

    def test_history_topic(self):
        assert classify_heading_topic("景区概况与千年历史渊源") == "history"
        assert classify_heading_topic("祥符禅寺的千年兴衰") == "history"

    def test_culture_topic(self):
        assert classify_heading_topic("核心文化内涵") == "culture"

    def test_route_topic(self):
        assert classify_heading_topic("个性化游览路线推荐") == "route"

    def test_general_topic(self):
        assert classify_heading_topic("实用游览贴士") == "general"


class TestGuideParsing:
    """Test guide DOCX section parsing."""

    def test_parse_guide_sections(self, tmp_path):
        doc = Document()
        doc.add_heading("景区概况与千年历史渊源", level=1)
        doc.add_paragraph("灵山胜境历史可追溯至唐代。")
        doc.add_heading("小灵山的佛教缘起", level=2)
        doc.add_paragraph("玄奘法师途经马山，命名小灵山。")
        doc.add_heading("核心文化内涵", level=1)
        doc.add_paragraph("佛教文化深度传承。")

        path = tmp_path / "guide.docx"
        doc.save(path)

        loaded = Document(str(path))
        sections = parse_guide_sections(loaded)

        assert len(sections) == 2
        assert sections[0]["topic"] == "history"
        assert "小灵山的佛教缘起" in sections[0]["text"]
        assert sections[1]["topic"] == "culture"


class TestSpotTableParsing:
    """Test spots dataset table parsing."""

    def test_parse_spots_tables(self, tmp_path):
        doc = Document()
        table = doc.add_table(rows=2, cols=3)
        table.rows[0].cells[0].text = "景区名称"
        table.rows[0].cells[1].text = "景点ID"
        table.rows[0].cells[2].text = "景点名称"
        table.rows[1].cells[0].text = "灵山胜境"
        table.rows[1].cells[1].text = "LS-001"
        table.rows[1].cells[2].text = "灵山大照壁"

        path = tmp_path / "spots.docx"
        doc.save(path)

        loaded = Document(str(path))
        records = parse_spots_tables(loaded)

        assert len(records) == 1
        assert records[0]["景点ID"] == "LS-001"
        assert records[0]["景点名称"] == "灵山大照壁"


class TestFAQExtraction:
    """Test FAQ extraction helpers."""

    def test_extract_ticket_faq(self):
        text = "成人票 | 210元；半价票 | 105元；免票 | 0元。"
        faqs = extract_ticket_faq(text)
        assert any(f["question"] == "灵山胜境成人票多少钱？" for f in faqs)

    def test_build_static_ticket_faqs(self):
        faqs = build_static_ticket_faqs()
        questions = {f["question"] for f in faqs}
        assert "灵山胜境门票多少钱？" in questions
        assert "灵山大佛有多高？" in questions

    def test_extract_route_faqs(self):
        sections = [
            {"topic": "route", "text": "【历史文化爱好者路线（6小时深度游）】\n路线规划：南门入园→灵山大佛→出口\n讲解重点：历史"},
        ]
        faqs = extract_route_faqs(sections)
        assert len(faqs) == 1
        assert faqs[0]["category"] == "route"

    def test_extract_spot_faqs(self):
        records = [
            {"景点名称": "灵山大佛", "建筑/景观参数": "高88米", "具体位置": "祥符禅寺北侧"},
        ]
        faqs = extract_spot_faqs(records)
        questions = {f["question"] for f in faqs}
        assert "灵山大佛有多高？" in questions
        assert "灵山大佛在哪里？" in questions


class TestImportFunctions:
    """Placeholder for DB import integration tests."""
    pass

